import os
from django.conf import settings
from reportlab.lib import enums
from reportlab.lib.pagesizes import A4
from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Paragraph, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.utils import ImageReader
from parkstay.models import Booking
import re
from docx import Document

def docx_replace_regex(doc_obj, regex , replace, key, bold_font=False,italic_font=False, underline_font=False):

    for p in doc_obj.paragraphs:
        #p.text = regex.sub(replace, p.text)
        style = p.style
        if key in p.text:
        #    print (p.text)
            p.bold = True
            p.text = regex.sub(replace, '')
            p.add_run(replace).bold = bold_font 
            #p.style.font.bold = True
        if regex.search(p.text):
            inline = p.runs
            #print (inline)
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            #print (row)
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace, key, bold_font, italic_font, underline_font)



def create_confirmation(booking):
    confirmation_doc = settings.BASE_DIR+"/parkstay/templates/doc/booking_confirmation_template.docx"
    temp_directory = settings.BASE_DIR+"/tmp/"

    doc = Document(confirmation_doc)


    if booking.first_campsite_list:
        campsites = []
        if booking.campground.site_type == 0:
            for item in booking.first_campsite_list:
                campsites.append(item.name if item else "")
        elif booking.campground.site_type == 1 or 2:
            for item in booking.first_campsite_list:
                campsites.append(item.type.split(':', 1)[0] if item else "")
        campsite = ', '.join(campsites)
        result = {x: campsites.count(x) for x in campsites}
        for key, value in result.items():
            campsite = ', '.join(['%sx %s' % (value, key) for (key, value) in result.items()])

    vehicle_data = ""
    if booking.vehicle_payment_status:
        for r in booking.vehicle_payment_status:
            data = r['Type']+"\t"+r['Rego']
            if r.get('Paid') is not None:
                if r['Paid'] == 'Yes':
                   data = data+'\tEntry fee paid'
                elif r['Paid'] == 'No':
                   data = data+'\tUnpaid'
                elif r['Paid'] == 'pass_required':
                   data = data+'\tPark Pass Required'
            vehicle_data = vehicle_data + data+"\n"

    key = "{{campname}}"
    docx_replace_regex(doc, re.compile(r''+key) , '{}, {}'.format(booking.campground.name, booking.campground.park.name), key, True)
    key = "{{campsite}}"
    docx_replace_regex(doc, re.compile(r''+key) , campsite, key, False)
    key = "{{bookingdates}}"
    docx_replace_regex(doc, re.compile(r''+key) , booking.stay_dates, key,False)
    key = "{{numberofguests}}"
    docx_replace_regex(doc, re.compile(r''+key) , booking.stay_guests, key, False)
    key = "{{nameandemail}}"
    docx_replace_regex(doc, re.compile(r''+key) , '{} {} ({})'.format(booking.details.get('first_name', ''), booking.details.get('last_name', ''), booking.customer.email if booking.customer else None), key,False)
    key = "{{bookingno}}"
    docx_replace_regex(doc, re.compile(r''+key) , booking.confirmation_number, key,False)
    key = "{{vehicles}}"
    docx_replace_regex(doc, re.compile(r''+key) , vehicle_data, key,False)
    key = "{{additionalinfo}}"
    docx_replace_regex(doc, re.compile(r''+key) , booking.campground.additional_info, key, False)
    
    try:
        os.stat(temp_directory)
    except:
        os.mkdir(temp_directory)
    new_doc_file = temp_directory+'booking_confirmation_'+str(booking.id)+'.docx'
    new_pdf_file = temp_directory+'booking_confirmation_'+str(booking.id)+'.pdf'
    doc.save(new_doc_file)
    os.system("libreoffice --headless --convert-to pdf "+new_doc_file+" --outdir "+temp_directory)
   
    confirmation_buffer = None
    with open(new_pdf_file, 'rb') as f:
         confirmation_buffer = f.read()
    os.remove(new_doc_file)
    os.remove(new_pdf_file)
    return confirmation_buffer

